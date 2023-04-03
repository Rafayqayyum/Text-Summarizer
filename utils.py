import openai
import docx
import PyPDF2
import os
max_words=2500
max_response=500

#function to generate summary
def generate_summary(prompt,api_key):
  openai.api_key = api_key
  prompt = prompt.encode(
    'ascii', errors='ignore'
  ).decode('ascii')
  words=prompt.split()
  prompt = " ".join(words[:max_words]).strip()
  tokens_discarded=len(words)-max_words
  if tokens_discarded<0:
    tokens_discarded=0
  try:
    response = openai.Completion.create(
      model="text-davinci-003",
      prompt=f"{prompt} tl;dr:",
      temperature=0.91,
      max_tokens=max_response,
      top_p=0.8,
      frequency_penalty=0.1,
      presence_penalty=0.1)
    return True,response.choices[0].text.strip(),tokens_discarded
  # check for openai api key error
  except openai.error.AuthenticationError as e:
    return False,'Invalid OpenAI API Key',0
  except openai.error.PermissionError as e:
    return False,'Your OpenAI API Key does not have access to this model',0
  except openai.error.RateLimitError as e:
    return False,'OpenAI API Key Rate Limit Reached',0
  except Exception as e:
    return False,None,0


#function to read docx file
def read_docx(filename):
  try:
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return True,' '.join(fullText).strip().replace('\n','')
  except:
    return False,0

# function to save docx file
def save_docx(text,filename):
  document = docx.Document()
  document.add_heading("Summary")
  document.add_paragraph(text)
  try:
    document.save(filename)
    document
    return True
  except:
    return False

# function to read pdf file
def read_pdf(filename):
  try:
    pdf = open(filename, 'rb')
    pdf_reader = PyPDF2.PdfReader(pdf)
    text=''
    for i in range(0,len(pdf_reader.pages)):
      page = pdf_reader.pages[i]
      text=text+ ' '+page.extract_text().strip().replace('\n', '')
    return True,text
  except:
    return False,0
